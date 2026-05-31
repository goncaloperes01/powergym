from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "Relatorio_PowerGym.docx"
ASSETS = BASE / "relatorio_assets"
ASSETS.mkdir(exist_ok=True)


def font(size=26, bold=False):
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline="#CBD5E1", radius=18, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fnt, fill="#111827"):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=4, align="center")
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2), text, font=fnt, fill=fill, spacing=4, align="center")


def draw_arrow(draw, p1, p2, label="", color="#334155"):
    draw.line([p1, p2], fill=color, width=3)
    if label:
        mx, my = (p1[0] + p2[0]) / 2, (p1[1] + p2[1]) / 2
        b = draw.textbbox((0, 0), label, font=font(18, True))
        w, h = b[2] - b[0], b[3] - b[1]
        draw.rounded_rectangle((mx - w / 2 - 8, my - h / 2 - 5, mx + w / 2 + 8, my + h / 2 + 5), radius=8, fill="#FFFFFF", outline="#CBD5E1")
        draw.text((mx - w / 2, my - h / 2), label, font=font(18, True), fill="#0F172A")


def make_er_diagram(path):
    img = Image.new("RGB", (2000, 1350), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    title_font = font(44, True)
    entity_font = font(26, True)
    small_font = font(18)
    draw.text((70, 45), "Modelo Entidade-Relacionamento - PowerGym", font=title_font, fill="#0F172A")
    draw.text((70, 100), "Entidades principais e cardinalidades funcionais do sistema.", font=font(22), fill="#64748B")

    boxes = {
        "User": (90, 200, 380, 340),
        "Socio": (500, 185, 820, 365),
        "Treinador": (910, 185, 1230, 365),
        "Modalidade": (1320, 200, 1640, 340),
        "Aulas": (500, 460, 820, 650),
        "Inscricao": (910, 460, 1230, 650),
        "PlanoTreino": (1320, 460, 1640, 650),
        "Exercicio": (500, 745, 820, 915),
        "Pagamento": (910, 745, 1230, 915),
    }
    descriptions = {
        "User": "Conta Django\nlogin e permissões",
        "Socio": "Membro do ginásio\nperfil, foto, dados",
        "Treinador": "Profissional\nespecialidade e aulas",
        "Modalidade": "Tipo de treino\nnome e intensidade",
        "Aulas": "Sessão marcada\ndata, hora, lotação",
        "Inscricao": "Sócio inscrito\nnuma aula",
        "PlanoTreino": "Objetivo atribuído\na um sócio",
        "Exercicio": "Exercícios do plano\nséries e repetições",
        "Pagamento": "Mensalidades\nestado financeiro",
    }
    for name, box in boxes.items():
        fill = "#FFFFFF"
        if name in {"Socio", "Treinador"}:
            fill = "#FFF7ED"
        elif name in {"Aulas", "Inscricao", "PlanoTreino"}:
            fill = "#EFF6FF"
        elif name in {"Pagamento"}:
            fill = "#F0FDF4"
        rounded_box(draw, box, fill=fill, outline="#94A3B8")
        x1, y1, x2, y2 = box
        draw.text((x1 + 20, y1 + 18), name, font=entity_font, fill="#0F172A")
        draw.line((x1 + 20, y1 + 58, x2 - 20, y1 + 58), fill="#CBD5E1", width=2)
        draw.multiline_text((x1 + 20, y1 + 78), descriptions[name], font=small_font, fill="#475569", spacing=5)

    draw.text((70, 1015), "Relações e cardinalidades", font=font(30, True), fill="#0F172A")
    relationships = [
        ("auth_user", "Socio", "1 : 0..1", "Conta opcional associada ao perfil de sócio"),
        ("auth_user", "Treinador", "1 : 0..1", "Conta opcional associada ao perfil de treinador"),
        ("Socio", "Pagamento", "1 : N", "Um sócio pode ter vários pagamentos"),
        ("Socio", "Inscricao", "1 : N", "Um sócio pode ter várias inscrições"),
        ("Aulas", "Inscricao", "1 : N", "Uma aula pode receber várias inscrições"),
        ("Modalidade", "Aulas", "1 : N", "Uma modalidade pode originar várias aulas"),
        ("Treinador", "Aulas", "1 : N", "Um treinador pode lecionar várias aulas"),
        ("Socio", "PlanoTreino", "1 : N", "Um sócio pode ter vários planos ao longo do tempo"),
        ("Treinador", "PlanoTreino", "1 : N", "Um treinador pode criar vários planos"),
        ("PlanoTreino", "Exercicio", "1 : N", "Um plano contém vários exercícios"),
    ]
    col_x = [70, 430, 790, 1150]
    row_y = 1070
    for idx, rel in enumerate(relationships):
        x = col_x[idx % 4]
        y = row_y + (idx // 4) * 82
        draw.rounded_rectangle((x, y, x + 330, y + 64), radius=12, fill="#FFFFFF", outline="#CBD5E1")
        draw.text((x + 14, y + 10), f"{rel[0]} -> {rel[1]}", font=font(17, True), fill="#0F172A")
        draw.text((x + 14, y + 34), f"{rel[2]} | {rel[3]}", font=font(13), fill="#475569")

    draw.text((70, 1290), "Nota: auth_user é a tabela de utilizadores do Django; Socio e Treinador podem estar associados a contas de autenticação.", font=font(18), fill="#64748B")
    img.save(path)


def make_physical_diagram(path):
    img = Image.new("RGB", (1900, 1500), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 40), "Modelo Físico da Base de Dados", font=font(44, True), fill="#0F172A")
    draw.text((70, 96), "Tabelas, chaves primárias e chaves estrangeiras implementadas em Django/PostgreSQL.", font=font(22), fill="#64748B")

    tables = {
        "auth_user": ["id PK", "username", "email", "password", "is_superuser", "is_staff"],
        "ginasio_socio": ["id PK", "user_id FK -> auth_user.id", "nome", "email UNIQUE", "contacto", "dataNascimento", "dataAdesao", "foto"],
        "ginasio_treinador": ["id PK", "user_id FK -> auth_user.id", "nome", "especialidade", "biografia", "foto"],
        "ginasio_modalidade": ["id PK", "nome", "descricao", "intensidade"],
        "ginasio_aulas": ["id PK", "modalidade_id FK", "treinador_id FK", "data", "hora", "lotacao"],
        "ginasio_inscricao": ["id PK", "socio_id FK", "aula_id FK", "dataRegisto"],
        "ginasio_planotreino": ["id PK", "socio_id FK", "treinador_id FK", "objetivo", "data_inicio"],
        "ginasio_exercicio": ["id PK", "planoTreino_id FK", "nome", "series", "repeticoes", "descanso_segundos"],
        "ginasio_pagamento": ["id PK", "socio_id FK", "valor", "data_vencimento", "data_pago", "estado", "frequencia", "descritivo"],
    }
    positions = {
        "auth_user": (70, 170),
        "ginasio_socio": (430, 160),
        "ginasio_treinador": (850, 160),
        "ginasio_modalidade": (1270, 170),
        "ginasio_pagamento": (70, 590),
        "ginasio_inscricao": (430, 590),
        "ginasio_aulas": (850, 560),
        "ginasio_planotreino": (1270, 590),
        "ginasio_exercicio": (1270, 1010),
    }
    widths = {k: 340 for k in tables}
    for name, fields in tables.items():
        x, y = positions[name]
        h = 70 + len(fields) * 34
        rounded_box(draw, (x, y, x + widths[name], y + h), fill="#F8FAFC", outline="#94A3B8", radius=12)
        draw.rounded_rectangle((x, y, x + widths[name], y + 48), radius=12, fill="#1F2937", outline="#1F2937")
        draw.text((x + 16, y + 12), name, font=font(20, True), fill="#FFFFFF")
        for i, field in enumerate(fields):
            fill = "#0F172A" if "PK" in field else "#334155"
            if "FK" in field:
                fill = "#1D4ED8"
            draw.text((x + 16, y + 62 + i * 34), field, font=font(17), fill=fill)

    legend_y = 1310
    draw.rounded_rectangle((70, legend_y, 1830, legend_y + 120), radius=14, fill="#F8FAFC", outline="#CBD5E1")
    draw.text((95, legend_y + 18), "Relações físicas principais", font=font(22, True), fill="#0F172A")
    rel_text = (
        "auth_user 1:0..1 ginasio_socio | auth_user 1:0..1 ginasio_treinador | "
        "ginasio_socio 1:N ginasio_pagamento, ginasio_inscricao e ginasio_planotreino | "
        "ginasio_treinador 1:N ginasio_aulas e ginasio_planotreino | "
        "ginasio_modalidade 1:N ginasio_aulas | ginasio_aulas 1:N ginasio_inscricao | "
        "ginasio_planotreino 1:N ginasio_exercicio"
    )
    draw.multiline_text((95, legend_y + 55), rel_text, font=font(17), fill="#475569", spacing=6)

    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True, "FFFFFF")
        set_cell_shading(hdr[i], "1F2937")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    r.font.size = Pt(10.5)
    p.add_run("\n" + text)
    doc.add_paragraph()


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "PowerGym - Relatório do Projeto"
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(100, 116, 139)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.text = "Bases de Dados / Django"
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def main():
    er_path = ASSETS / "diagrama_er_powergym.png"
    physical_path = ASSETS / "modelo_fisico_powergym.png"
    make_er_diagram(er_path)
    make_physical_diagram(physical_path)

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PowerGym")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(17, 24, 39)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Relatório do Projeto de Bases de Dados e Aplicação Web em Django")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(71, 85, 105)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Curso: Bases de Dados / Bases de Dados e Análise de Informação\n").bold = True
    meta.add_run("Contexto: Sistema de gestão de ginásio com landing page pública, dashboard de administração, área de sócio e área de treinador\n")
    meta.add_run("Domínio: powergym.site | Deploy: Render | Base de dados: PostgreSQL/Neon")
    doc.add_paragraph()
    add_callout(doc, "Resumo executivo", "O PowerGym é uma aplicação web desenvolvida em Django para gerir o funcionamento de um ginásio. O sistema integra uma componente pública de apresentação, autenticação de utilizadores, gestão administrativa, perfis de sócio e treinador, aulas, inscrições, planos de treino, exercícios e pagamentos. O projeto foi desenhado para cumprir os requisitos do enunciado, nomeadamente uma base de dados relacional com mais de cinco entidades, relações 1:N, formulários próprios, templates Django, administração via Django Admin, menus de navegação e uma página index/dashboard com contagens das principais tabelas.")
    doc.add_page_break()

    doc.add_heading("Índice", level=1)
    contents = [
        "1. Introdução e enquadramento",
        "2. Objetivos do projeto",
        "3. Requisitos funcionais",
        "4. Arquitetura da aplicação",
        "5. Modelo de dados e decisões de base de dados",
        "6. Diagramas ER e físico",
        "7. Funcionalidades implementadas",
        "8. Segurança, autenticação e permissões",
        "9. Interface, usabilidade e landing page",
        "10. Deploy, base de dados em produção e ficheiros estáticos",
        "11. Testes e validação",
        "12. Limitações e trabalho futuro",
        "13. Distribuição de trabalho e expectativa de nota",
        "14. Conclusão",
    ]
    add_numbered(doc, contents)
    doc.add_page_break()

    doc.add_heading("1. Introdução e enquadramento", level=1)
    doc.add_paragraph("O presente relatório descreve o desenvolvimento do PowerGym, uma aplicação web para gestão de ginásio construída com Django e base de dados relacional. A escolha deste contexto foi motivada por ser um cenário realista e suficientemente rico para demonstrar modelação de dados, relações entre entidades, operações CRUD, autenticação, permissões, formulários, templates e integração com serviços externos.")
    doc.add_paragraph("A aplicação foi evoluindo desde uma estrutura inicial de CRUDs simples até uma solução mais completa: primeiro foram criados os modelos centrais e as páginas administrativas; depois foram acrescentadas funcionalidades específicas do domínio, como inscrição de sócios em aulas, criação de planos de treino, gestão de pagamentos e perfis personalizados. Na fase final, o projeto foi preparado para produção com deploy no Render, base de dados PostgreSQL no Neon, armazenamento de imagens com Cloudinary e domínio próprio powergym.site.")

    doc.add_heading("2. Objetivos do projeto", level=1)
    add_bullets(doc, [
        "Criar uma aplicação Django funcional e executável, com código organizado por modelos, views, forms, urls e templates.",
        "Construir uma base de dados relacional com pelo menos cinco entidades próprias, excluindo tabelas de ligação, e com relações 1:N.",
        "Disponibilizar CRUDs completos para as entidades principais através de páginas desenvolvidas pelo grupo.",
        "Permitir a gestão das mesmas entidades também através do Django Admin.",
        "Criar uma página index/dashboard com ligações e contagens das principais tabelas.",
        "Evoluir a aplicação para um produto mais profissional, com landing page pública, áreas autenticadas e design mais cuidado.",
        "Colocar a aplicação online com base de dados PostgreSQL, domínio próprio e armazenamento externo para imagens.",
    ])

    doc.add_heading("3. Requisitos funcionais", level=1)
    requirements = [
        ("RF01", "Landing page pública", "O visitante deve conseguir aceder a uma página inicial pública com apresentação do PowerGym, modalidades em destaque, equipa e botões para login ou registo."),
        ("RF02", "Registo de sócio", "Um visitante deve conseguir criar uma conta de sócio, indicando dados pessoais, credenciais e opcionalmente uma fotografia de perfil."),
        ("RF03", "Autenticação", "O sistema deve permitir login e logout, redirecionando utilizadores para a área adequada conforme o tipo de conta."),
        ("RF04", "Dashboard de administração", "O administrador deve aceder a um painel com estatísticas do ginásio, incluindo contagens de sócios, treinadores, aulas, modalidades, inscrições, planos, exercícios e pagamentos."),
        ("RF05", "Gestão de sócios", "O administrador deve poder listar, consultar, criar, editar e apagar sócios."),
        ("RF06", "Gestão de treinadores", "O administrador deve poder gerir treinadores, especialidades, biografias, fotografias e contas associadas."),
        ("RF07", "Gestão de modalidades e aulas", "O administrador deve poder gerir modalidades e criar aulas com data, hora, lotação, treinador e modalidade."),
        ("RF08", "Inscrição em aulas", "O sistema deve permitir inscrever sócios em aulas, evitando inscrições duplicadas na mesma aula."),
        ("RF09", "Área de sócio", "O sócio deve consultar dados pessoais, passe digital por QR Code, aulas inscritas, pagamentos e plano de treino."),
        ("RF10", "Área de treinador", "O treinador deve consultar as aulas associadas ao seu perfil."),
        ("RF11", "Planos de treino e exercícios", "O administrador deve criar planos de treino por sócio e treinador, e associar exercícios com séries, repetições e descanso."),
        ("RF12", "Gestão de pagamentos", "O administrador deve criar e controlar pagamentos, estados pendentes, liquidados e atrasados, incluindo mensalidades."),
        ("RF13", "Gestão de acessos", "O administrador deve criar contas de administrador, treinador ou associar contas a sócios/treinadores existentes."),
        ("RF14", "Proteção de páginas", "As páginas administrativas devem estar protegidas, sendo acessíveis apenas a administradores autenticados."),
    ]
    add_table(doc, ["ID", "Requisito", "Descrição"], requirements, widths=[1.2, 4.0, 10.5])

    doc.add_heading("4. Arquitetura da aplicação", level=1)
    doc.add_paragraph("A aplicação segue a arquitetura habitual de um projeto Django. O projeto principal chama-se ginasio_project e contém as configurações, urls globais e ficheiros WSGI/ASGI. A aplicação de domínio chama-se ginasio e concentra a lógica do ginásio: modelos, formulários, views, urls, templates e configuração do admin.")
    add_table(doc, ["Componente", "Ficheiros principais", "Responsabilidade"], [
        ("Projeto Django", "ginasio_project/settings.py, urls.py, wsgi.py", "Configuração global, base de dados, static/media, middleware, deploy e roteamento principal."),
        ("Aplicação ginasio", "models.py, views.py, forms.py, urls.py", "Implementação do domínio PowerGym e das regras funcionais."),
        ("Templates", "ginasio/templates/ginasio/*.html", "Interface HTML das páginas públicas, administrativas e perfis."),
        ("Estáticos", "static/css/style.css, static/images/*", "Estilos globais, logo, favicons e recursos visuais."),
        ("Admin", "ginasio/admin.py", "Registo das entidades no Django Admin com filtros e pesquisa."),
        ("Migrations", "ginasio/migrations/*.py", "Histórico de evolução do modelo físico da base de dados."),
    ], widths=[3.5, 5.2, 7.0])

    doc.add_heading("5. Modelo de dados e decisões de base de dados", level=1)
    doc.add_paragraph("O modelo de dados foi desenhado para representar o funcionamento de um ginásio real. Em vez de limitar a aplicação a uma lista isolada de membros, foram criadas entidades ligadas entre si: sócios participam em aulas, aulas pertencem a modalidades e são dadas por treinadores, sócios têm pagamentos e planos de treino, e planos de treino contêm exercícios.")
    doc.add_paragraph("A base de dados em produção usa PostgreSQL através do Neon. Durante o desenvolvimento local foi usado SQLite como fallback, mas a configuração final lê DATABASE_URL através de variáveis de ambiente, permitindo que o Render se ligue à base de dados Neon de forma segura.")
    add_table(doc, ["Entidade", "Finalidade", "Relações principais"], [
        ("Socio", "Representa um membro do ginásio.", "Pode ter vários pagamentos, inscrições e planos de treino. Pode estar associado a auth_user."),
        ("Treinador", "Representa um profissional do ginásio.", "Pode lecionar várias aulas e criar vários planos de treino. Pode estar associado a auth_user."),
        ("Modalidade", "Representa um tipo de treino, como musculação, funcional ou natação.", "Pode estar associada a várias aulas."),
        ("Aulas", "Representa uma sessão marcada no calendário.", "Pertence a uma modalidade e a um treinador; pode ter várias inscrições."),
        ("Inscricao", "Liga um sócio a uma aula.", "É a entidade associativa entre Socio e Aulas, com data de registo."),
        ("PlanoTreino", "Representa o plano atribuído a um sócio.", "Pertence a um sócio e a um treinador; contém vários exercícios."),
        ("Exercicio", "Representa um exercício dentro de um plano.", "Pertence a um plano de treino."),
        ("Pagamento", "Representa mensalidades ou pagamentos únicos.", "Pertence a um sócio e guarda valor, datas, estado e frequência."),
    ], widths=[3.0, 5.2, 7.5])

    doc.add_heading("6. Diagramas ER e físico", level=1)
    doc.add_paragraph("O diagrama ER apresenta a visão conceptual das entidades e relações. O modelo físico mostra a tradução para tabelas e campos concretos usados pela aplicação Django e pela base de dados PostgreSQL.")
    doc.add_picture(str(er_path), width=Inches(6.7))
    p = doc.add_paragraph("Figura 1 - Diagrama entidade-relacionamento do PowerGym.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_picture(str(physical_path), width=Inches(6.7))
    p = doc.add_paragraph("Figura 2 - Modelo físico simplificado das tabelas principais.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("7. Funcionalidades implementadas", level=1)
    doc.add_heading("7.1 Landing page pública", level=2)
    doc.add_paragraph("A landing page foi reformulada para funcionar como entrada pública do PowerGym. Mostra o nome da marca, uma mensagem de valor, botões para começar ou fazer login, estatísticas vindas da base de dados, modalidades, equipa e uma secção de chamada para a área de sócio. A página usa dados reais das tabelas Modalidade, Treinador, Socio e Aulas.")
    doc.add_heading("7.2 Dashboard de administração", level=2)
    doc.add_paragraph("O dashboard resume o estado geral do ginásio. Para além das contagens principais, inclui gráficos com Chart.js: um gráfico circular para estados financeiros e um gráfico de barras para volume de dados. Esta página cumpre o requisito de existir uma index/dashboard com links e contagens das tabelas.")
    doc.add_heading("7.3 CRUDs administrativos", level=2)
    doc.add_paragraph("Foram implementadas páginas de listagem, detalhe, criação, edição e eliminação para as entidades principais. Os formulários usam ModelForm, reduzindo duplicação entre o modelo e a interface. As páginas de confirmação de eliminação reduzem o risco de apagar dados por engano.")
    doc.add_heading("7.4 Área de sócio", level=2)
    doc.add_paragraph("Após login, um sócio acede à sua área pessoal. Esta página mostra dados pessoais, QR Code de acesso, aulas em que está inscrito, pagamentos e plano de treino. O QR Code é gerado dinamicamente com a biblioteca qrcode e convertido para Base64 para ser apresentado diretamente no template.")
    doc.add_heading("7.5 Área de treinador", level=2)
    doc.add_paragraph("A área de treinador permite que um profissional veja as aulas associadas ao seu perfil. A associação é feita através do campo user em Treinador, permitindo distinguir contas de sócio, treinador e administrador.")
    doc.add_heading("7.6 Pagamentos e mensalidades", level=2)
    doc.add_paragraph("O sistema gere pagamentos com estados Liquidado, Pendente e Atrasado. As views atualizam pagamentos pendentes para Atrasado quando a data de vencimento já passou. Quando uma mensalidade é liquidada, a lógica cria a mensalidade seguinte, automatizando parte do processo financeiro.")

    doc.add_heading("8. Segurança, autenticação e permissões", level=1)
    doc.add_paragraph("A autenticação usa o sistema padrão do Django. A view redirecionar_login decide para onde cada utilizador deve ir depois de entrar: administrador para o dashboard, treinador para a área de treinador e sócio para a área de cliente.")
    add_bullets(doc, [
        "As páginas administrativas usam um decorador admin_required, que bloqueia acesso a visitantes, sócios e treinadores.",
        "O sócio pode cancelar apenas as suas próprias inscrições.",
        "As credenciais de produção foram removidas do código e passaram a variáveis de ambiente no Render.",
        "DEBUG deve estar False em produção.",
        "SECRET_KEY, DATABASE_URL e credenciais Cloudinary são configuradas fora do repositório.",
    ])

    doc.add_heading("9. Interface, usabilidade e landing page", level=1)
    doc.add_paragraph("A interface foi melhorada em duas fases. Primeiro, a landing page foi redesenhada para causar boa primeira impressão e explicar rapidamente a proposta do PowerGym. Depois, as páginas internas foram harmonizadas com uma navbar escura, botões consistentes, cards, tabelas e formulários com aspeto mais profissional.")
    doc.add_paragraph("Também foi acrescentado feedback de erro no login. Quando um utilizador tenta entrar com nome ou palavra-passe incorretos, aparece uma mensagem clara em vermelho. Esta decisão melhora a experiência do utilizador porque evita uma falha silenciosa.")

    doc.add_heading("10. Deploy, base de dados em produção e ficheiros estáticos", level=1)
    doc.add_paragraph("O projeto foi colocado online no Render e ligado ao domínio powergym.site. A base de dados de produção encontra-se no Neon, usando PostgreSQL. As imagens enviadas pelos utilizadores são guardadas no Cloudinary através de django-cloudinary-storage. Os ficheiros estáticos, como CSS, logo e favicons, são servidos com apoio de WhiteNoise.")
    add_table(doc, ["Serviço", "Uso no projeto", "Justificação"], [
        ("Render", "Alojamento da aplicação Django.", "Permite deploy automático a partir do GitHub e execução com Gunicorn."),
        ("Neon", "Base de dados PostgreSQL.", "Cumpre o requisito de usar PostgreSQL e separa os dados do servidor da aplicação."),
        ("Cloudinary", "Armazenamento de imagens de sócios/treinadores.", "Evita depender do disco efémero do Render para ficheiros enviados."),
        ("WhiteNoise", "Serviço de ficheiros estáticos.", "Permite servir CSS, logo e recursos estáticos em produção."),
        ("GitHub", "Controlo de versões e integração com deploy.", "Regista evolução do projeto e facilita publicação no Render."),
    ], widths=[2.5, 5.2, 8.0])

    doc.add_heading("11. Testes e validação", level=1)
    doc.add_paragraph("Foram realizados testes manuais e verificações automáticas simples durante a fase final. O objetivo foi garantir que as páginas principais carregavam, que o login inválido apresentava mensagem, que o dashboard funcionava, que as páginas administrativas estavam protegidas e que não existiam migrations pendentes.")
    add_bullets(doc, [
        "python manage.py check sem erros.",
        "python manage.py makemigrations --check --dry-run sem alterações pendentes.",
        "Validação de landing page, login, dashboard e rotas principais.",
        "Correção de erro 500 na página de apagar plano de treino.",
        "Correção de permissões para impedir acesso público às páginas de gestão.",
        "Correção de ficheiros estáticos e logo em produção.",
    ])

    doc.add_heading("12. Limitações e trabalho futuro", level=1)
    add_bullets(doc, [
        "Criar testes automatizados permanentes para views e formulários.",
        "Adicionar validação de lotação das aulas, impedindo inscrições acima da capacidade.",
        "Melhorar relatórios financeiros com filtros por mês e exportação.",
        "Permitir que treinadores editem planos dos seus próprios sócios, caso se pretenda esse fluxo.",
        "Adicionar recuperação de password por email.",
        "Criar histórico de check-ins através do QR Code.",
    ])

    doc.add_heading("13. Distribuição de trabalho e expectativa de nota", level=1)
    doc.add_paragraph("Esta secção deve ser preenchida pelo grupo antes da entrega, conforme pedido no enunciado. A tabela abaixo foi deixada preparada para indicar percentagem de trabalho e expectativa de nota individual.")
    add_table(doc, ["Elemento", "% de trabalho", "Principais contributos", "Expectativa de nota"], [
        ("Elemento 1", "A preencher", "A preencher", "A preencher"),
        ("Elemento 2", "A preencher", "A preencher", "A preencher"),
    ], widths=[3.0, 2.5, 7.2, 3.0])

    doc.add_heading("14. Conclusão", level=1)
    doc.add_paragraph("O PowerGym cumpre os objetivos centrais do projeto: é uma aplicação Django funcional, com base de dados relacional, mais de cinco entidades relacionadas, relações 1:N, CRUDs próprios, Django Admin, dashboard com contagens, autenticação, permissões e deploy em produção. Para além dos requisitos mínimos, a aplicação foi evoluída para uma experiência mais profissional, com landing page pública, áreas diferenciadas por tipo de utilizador, gráficos, QR Code, Cloudinary, Neon, Render e domínio próprio.")
    doc.add_paragraph("A modelação da base de dados foi pensada para representar um cenário real de ginásio e não apenas um exemplo simplificado. As relações entre sócios, aulas, inscrições, treinadores, planos, exercícios e pagamentos demonstram o uso prático de chaves estrangeiras, integridade relacional e lógica de domínio.")

    doc.add_page_break()
    doc.add_heading("Apêndice A - Mapeamento de entidades para ficheiros Django", level=1)
    add_table(doc, ["Entidade", "Model", "Form", "Views/URLs", "Templates"], [
        ("Socio", "Socio", "SocioForm", "socio_list/detail/create/update/delete", "socio_*.html"),
        ("Treinador", "Treinador", "TreinadorForm", "treinador_list/detail/create/update/delete", "treinador_*.html"),
        ("Modalidade", "Modalidade", "ModalidadeForm", "modalidade_list/detail/create/update/delete", "modalidade_*.html"),
        ("Aulas", "Aulas", "AulasForm", "aula_list/detail/create/update/delete", "aula_*.html"),
        ("Inscricao", "Inscricao", "InscricaoForm", "inscricao_list/detail/create/update/delete", "inscricao_*.html"),
        ("PlanoTreino", "PlanoTreino", "PlanoTreinoForm", "planotreino_list/detail/create/update/delete", "planotreino_*.html"),
        ("Exercicio", "Exercicio", "ExercicioForm", "exercicio_list/detail/create/update/delete", "exercicio_*.html"),
        ("Pagamento", "Pagamento", "PagamentoForm", "pagamento_list/detail/create/update/delete", "pagamento_*.html"),
    ], widths=[2.7, 2.7, 3.0, 5.0, 3.7])

    doc.add_heading("Apêndice B - Decisões técnicas relevantes", level=1)
    add_bullets(doc, [
        "Django foi usado porque o enunciado e as aulas laboratoriais se baseiam neste framework.",
        "ModelForm foi usado para reduzir duplicação e manter formulários ligados diretamente aos modelos.",
        "PostgreSQL/Neon foi usado em produção para cumprir o requisito de base de dados PostgreSQL.",
        "Cloudinary foi usado porque o Render não garante persistência local para ficheiros enviados.",
        "WhiteNoise foi usado para servir ficheiros estáticos em produção.",
        "Gunicorn foi usado como servidor WSGI em produção no Render.",
        "Variáveis de ambiente foram usadas para evitar expor segredos no código.",
    ])

    add_header_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
