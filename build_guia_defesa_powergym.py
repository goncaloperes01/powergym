from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Guia_Defesa_PowerGym.docx")

BLUE = RGBColor(37, 99, 235)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(71, 85, 105)
ORANGE = RGBColor(249, 115, 22)
HEADER = "E8EEF5"
CALLOUT = "FFF7ED"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "CBD5E1")
        borders.append(tag)
    table._tbl.tblPr.append(borders)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def h(doc, level, text):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = BLUE if level <= 2 else DARK
    return para


def p(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.12
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    para.add_run(text).font.size = Pt(10)


def number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(3)
    para.add_run(text).font.size = Pt(10)


def code(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.18)
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 41, 59)


def callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    cell = table.rows[0].cells[0]
    shade_cell(cell, CALLOUT)
    para = cell.paragraphs[0]
    r1 = para.add_run(title + ": ")
    r1.bold = True
    r1.font.color.rgb = ORANGE
    r1.font.size = Pt(10)
    r2 = para.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_borders(table)
    headers = table.rows[0].cells
    shade_cell(headers[0], HEADER)
    shade_cell(headers[1], HEADER)
    set_cell_text(headers[0], "Elemento", True, DARK)
    set_cell_text(headers[1], "Explicacao para a defesa", True, DARK)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    doc.add_paragraph()


def three_col_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_borders(table)
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], HEADER)
        set_cell_text(table.rows[0].cells[i], header, True, DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 17, BLUE),
        ("Heading 2", 13.5, BLUE),
        ("Heading 3", 11.5, DARK),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    return doc


def add_cover(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run("PowerGym")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = BLUE
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run("Guia de Defesa Tecnica do Projeto")
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = DARK
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run("Django + PostgreSQL/Neon + Render + Cloudinary")
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED
    doc.add_paragraph()
    callout(
        doc,
        "Objetivo",
        "Este documento foi feito para a defesa: explica o que cada ficheiro faz, como os pedidos passam por URLs, views, models e templates, como a base de dados Neon entra no sistema, como o Render faz host do site, e que decisoes tecnicas foram tomadas.",
    )
    kv_table(
        doc,
        [
            ("Projeto", "PowerGym - Sistema de gestao de ginasio"),
            ("Stack principal", "Django 6, PostgreSQL/Neon, Render, Cloudinary, WhiteNoise"),
            ("Data do guia", date.today().strftime("%d/%m/%Y")),
            ("Como estudar", "Le por secoes, depois treina os fluxos em voz alta: login, registo, CRUD, pagamentos, deploy."),
        ],
    )
    doc.add_page_break()


def add_content(doc):
    h(doc, 1, "1. Explicacao curta para abrir a defesa")
    p(
        doc,
        "O PowerGym e uma aplicacao web de gestao de ginasio desenvolvida em Django. Tem landing page publica, registo e login de socios, area de cliente, dashboard de administrador, gestao de socios, treinadores, modalidades, aulas, inscricoes, planos de treino, exercicios, pagamentos e acessos.",
    )
    p(
        doc,
        "A arquitetura segue o padrao MVT do Django: Models definem tabelas e relacoes, Views recebem pedidos e aplicam a logica, Templates mostram HTML, e URLs ligam enderecos do browser as views.",
    )
    callout(
        doc,
        "Frase segura",
        "O pedido entra pelo browser, o urls.py escolhe a view, a view consulta models/forms, envia dados no context e o template gera a pagina final.",
    )

    h(doc, 1, "2. Estrutura de pastas")
    kv_table(
        doc,
        [
            ("manage.py", "Comando principal do Django: runserver, check, test, makemigrations, migrate e createsuperuser."),
            ("ginasio_project/", "Configuracao global: settings.py, urls.py, wsgi.py e asgi.py."),
            ("ginasio/", "Aplicacao principal: models, views, forms, urls, admin e migrations."),
            ("ginasio/templates/ginasio/", "Templates HTML renderizados pelas views."),
            ("static/", "CSS, logo, favicon, manifest e imagens fixas."),
            ("media/", "Uploads locais em desenvolvimento; em producao as imagens vao para Cloudinary."),
            ("requirements.txt", "Dependencias instaladas pelo Render no deploy."),
        ],
    )

    h(doc, 1, "3. Arquitetura MVT")
    three_col_table(
        doc,
        ["Parte", "No projeto", "Responsabilidade"],
        [
            ("Model", "ginasio/models.py", "Define entidades e relacoes da base de dados."),
            ("View", "ginasio/views.py", "Recebe request, aplica regras, consulta/grava dados e devolve render ou redirect."),
            ("Template", "templates/ginasio/*.html", "Apresenta o HTML com variaveis vindas do context."),
            ("URLconf", "ginasio/urls.py", "Mapeia caminhos como /socios/ para views concretas."),
            ("Settings", "ginasio_project/settings.py", "Configuracao de apps, base de dados, static files, seguranca e deploy."),
        ],
    )

    h(doc, 1, "4. settings.py explicado")
    h(doc, 2, "4.1 SECRET_KEY, DEBUG e dominios")
    code(doc, "SECRET_KEY = os.environ.get('SECRET_KEY', 'django-insecure-dev-only-change-me')\nDEBUG = os.environ.get('DEBUG', 'False').lower() in ('1', 'true', 'yes', 'on')")
    p(doc, "A SECRET_KEY e lida das variaveis de ambiente do Render. Nao deve ficar no GitHub. DEBUG deve estar False em producao para nao mostrar erros tecnicos ao publico.")
    code(doc, "ALLOWED_HOSTS = 'localhost,127.0.0.1,testserver,powergym.site,www.powergym.site,.onrender.com'")
    p(doc, "ALLOWED_HOSTS define que dominios podem servir o site. Inclui localhost para desenvolvimento, o dominio powergym.site, www.powergym.site e o dominio tecnico do Render.")
    code(doc, "CSRF_TRUSTED_ORIGINS = 'https://powergym.site,https://www.powergym.site'")
    p(doc, "CSRF_TRUSTED_ORIGINS permite que formularios POST enviados pelo dominio HTTPS sejam aceites pelo Django.")

    h(doc, 2, "4.2 HTTPS e cookies")
    code(doc, "SECURE_PROXY_SSL_HEADER = ('HTTP_X_FORWARDED_PROTO', 'https')\nSECURE_SSL_REDIRECT = os.environ.get('SECURE_SSL_REDIRECT', str(not DEBUG)).lower() in (...)\nSESSION_COOKIE_SECURE = os.environ.get('SESSION_COOKIE_SECURE', str(not DEBUG)).lower() in (...)")
    p(doc, "O Render esta atras de um proxy. O Django precisa saber que o pedido original veio por HTTPS. As cookies de sessao e CSRF ficam seguras em producao.")

    h(doc, 2, "4.3 INSTALLED_APPS e MIDDLEWARE")
    kv_table(
        doc,
        [
            ("django.contrib.auth", "Sistema de utilizadores, login, logout, hashes de password e permissoes."),
            ("django.contrib.sessions", "Mantem o utilizador autenticado entre pedidos."),
            ("django.contrib.messages", "Mostra mensagens de sucesso, erro e aviso nos templates."),
            ("cloudinary_storage/cloudinary", "Uploads de imagens para Cloudinary."),
            ("django.contrib.staticfiles", "Gestao de CSS, favicon e imagens estaticas."),
            ("ginasio", "A aplicacao criada para o projeto."),
            ("WhiteNoiseMiddleware", "Permite servir static files em producao pelo Django/Render."),
            ("CsrfViewMiddleware", "Protege formularios contra ataques CSRF."),
            ("AuthenticationMiddleware", "Coloca request.user disponivel nas views/templates."),
        ],
    )

    h(doc, 2, "4.4 Base de dados: SQLite local e Neon em producao")
    code(doc, "DATABASES = {\n    'default': dj_database_url.config(\n        default=f\"sqlite:///{BASE_DIR / 'db.sqlite3'}\",\n        conn_max_age=600\n    )\n}")
    p(doc, "O dj_database_url le automaticamente a variavel DATABASE_URL. Se existir no Render, liga ao PostgreSQL da Neon. Se nao existir, usa db.sqlite3 local. Isto permite desenvolvimento local simples e producao com PostgreSQL.")
    callout(doc, "Como explicar Neon", "No Render colocamos DATABASE_URL com a connection string da Neon. O codigo nao contem o user/password da base de dados; le esses dados de variaveis de ambiente.")

    h(doc, 2, "4.5 Static files, media e Cloudinary")
    code(doc, "STATIC_URL = '/static/'\nSTATICFILES_DIRS = [os.path.join(BASE_DIR, 'static')]\nSTATIC_ROOT = os.path.join(BASE_DIR, 'staticfiles')")
    p(doc, "STATICFILES_DIRS e onde estao os ficheiros estaticos no codigo. STATIC_ROOT e para onde o collectstatic copia tudo no deploy.")
    code(doc, "STORAGES = {\n  'default': {'BACKEND': 'cloudinary_storage.storage.MediaCloudinaryStorage'},\n  'staticfiles': {'BACKEND': 'django.contrib.staticfiles.storage.StaticFilesStorage'}\n}")
    p(doc, "O storage default e Cloudinary, logo os ImageField de socios e treinadores guardam fotos fora do Render. Isto evita perder uploads quando o servidor reinicia.")

    h(doc, 1, "5. urls.py explicado")
    code(doc, "urlpatterns = [\n    path('favicon.ico', RedirectView.as_view(url=settings.STATIC_URL + 'images/favicon.ico', permanent=True)),\n    path('admin/', admin.site.urls),\n    path('', include('ginasio.urls')),\n]")
    p(doc, "ginasio_project/urls.py e o mapa principal. Trata /admin/, redireciona /favicon.ico para o favicon real e inclui todas as rotas da app ginasio.")
    p(doc, "ginasio/urls.py liga caminhos como /dashboard/, /socios/, /aulas/novo/ ou /pagamentos/<int:pk>/editar/ as respetivas views.")
    callout(doc, "O que e <int:pk>", "E o id do registo na base de dados. /socios/5/ chama socio_detail(request, pk=5).")

    h(doc, 1, "6. models.py: entidades e relacoes")
    three_col_table(
        doc,
        ["Model", "Campos principais", "Explicacao"],
        [
            ("Socio", "user, nome, email, contacto, dataNascimento, dataAdesao, foto", "Ficha de cliente do ginasio; pode estar ligada a um User para login."),
            ("Treinador", "user, nome, especialidade, biografia, foto", "Ficha de treinador; pode estar ligada a um User para dashboard proprio."),
            ("Modalidade", "nome, descricao, intensidade", "Tipo de atividade: Yoga, Natacao, Pilates, etc."),
            ("Aulas", "modalidade, treinador, data, hora, lotacao", "Aula concreta num dia e hora, dada por um treinador."),
            ("Inscricao", "socio, aula, dataRegisto", "Liga socios a aulas; funciona como tabela associativa."),
            ("PlanoTreino", "socio, treinador, objetivo, data_inicio", "Plano prescrito por treinador a um socio."),
            ("Exercicio", "planoTreino, nome, series, repeticoes, descanso_segundos", "Exercicios pertencentes a um plano."),
            ("Pagamento", "socio, valor, data_vencimento, data_pago, estado, frequencia, descritivo", "Mensalidades e pagamentos unicos."),
        ],
    )
    h(doc, 2, "6.1 Relacoes")
    kv_table(
        doc,
        [
            ("User 1-1 Socio", "Permite que um socio tenha conta de login."),
            ("User 1-1 Treinador", "Permite que um treinador tenha conta e dashboard."),
            ("Modalidade 1-N Aulas", "Uma modalidade pode aparecer em muitas aulas."),
            ("Treinador 1-N Aulas", "Um treinador pode dar muitas aulas."),
            ("Socio N-N Aulas via Inscricao", "Um socio pode estar em varias aulas e uma aula pode ter varios socios."),
            ("Socio 1-N PlanoTreino", "Um socio pode ter varios planos ao longo do tempo."),
            ("PlanoTreino 1-N Exercicio", "Um plano tem varios exercicios."),
            ("Socio 1-N Pagamento", "Um socio tem varias faturas/mensalidades."),
        ],
    )
    code(doc, "Socio --< Inscricao >-- Aulas -- Modalidade\n                    |\n                    +-- Treinador\n\nSocio --< PlanoTreino --< Exercicio\nTreinador --< PlanoTreino\nSocio --< Pagamento\nUser --1:1-- Socio / Treinador")
    p(doc, "As ForeignKeys usam on_delete=models.CASCADE. Se um registo pai for apagado, os filhos associados tambem podem ser removidos, mantendo a integridade da base de dados.")

    h(doc, 1, "7. Migrations")
    p(doc, "As migrations sao o historico da estrutura da base de dados. Quando alteramos models.py, criamos migrations e depois aplicamos na base de dados.")
    kv_table(
        doc,
        [
            ("0001_initial.py", "Criou as entidades iniciais."),
            ("0002 a 0004", "Evolucao do model Pagamento: datas, vencimento, frequencia e descritivo."),
            ("0005", "Adicionou ligacao Socio -> User."),
            ("0006", "Ajustou campos do Socio."),
            ("0007", "Adicionou ligacao Treinador -> User."),
            ("0008", "Ajustou Modalidade e intensidade."),
        ],
    )
    code(doc, "python manage.py makemigrations\npython manage.py migrate")

    h(doc, 1, "8. forms.py")
    p(doc, "Os forms sao ModelForms. Um ModelForm cria automaticamente campos HTML com base no model, reduzindo repeticao.")
    kv_table(
        doc,
        [
            ("SocioForm", "Cria/edita socios; usa input date e file input para foto."),
            ("TreinadorForm", "Cria/edita treinadores; aceita fotografia."),
            ("ModalidadeForm", "Cria/edita modalidades."),
            ("AulasForm", "Cria/edita aulas; usa date e time inputs."),
            ("InscricaoForm", "Cria/edita inscricoes; dataRegisto e automatico."),
            ("PlanoTreinoForm", "Cria/edita planos."),
            ("ExercicioForm", "Cria/edita exercicios."),
            ("PagamentoForm", "Cria/edita pagamentos; tem labels e classes Bootstrap."),
        ],
    )
    callout(doc, "Padrao de formulario", "GET mostra o formulario. POST valida com form.is_valid(). Se estiver valido, form.save() grava na base de dados.")

    h(doc, 1, "9. views.py: a logica principal")
    h(doc, 2, "9.1 Imports importantes")
    code(doc, "render, redirect, get_object_or_404\nUser, login_required, login, messages\nSocio, Treinador, Aulas, Modalidade, Inscricao, PlanoTreino, Exercicio, Pagamento")
    p(doc, "render devolve HTML, redirect muda de rota, get_object_or_404 procura um objeto e devolve 404 se nao existir. User e o utilizador nativo do Django.")

    h(doc, 2, "9.2 admin_required")
    code(doc, "if not request.user.is_authenticated:\n    return redirect('login')\nif not request.user.is_superuser:\n    return redirect('redirecionar_login')\nreturn view_func(request, *args, **kwargs)")
    p(doc, "Este decorator protege paginas administrativas. So utilizadores autenticados e superusers podem entrar nos CRUDs.")

    h(doc, 2, "9.3 landing_page")
    code(doc, "modalidades_destaque = Modalidade.objects.all()\nequipa_destaque = Treinador.objects.all()\ncontext = {'modalidades': ..., 'treinadores': ..., 'total_socios': ...}")
    p(doc, "A landing e publica e dinamica. Vai buscar modalidades, treinadores e contagens reais da base de dados para mostrar no site.")

    h(doc, 2, "9.4 index dashboard")
    code(doc, "Pagamento.objects.filter(estado='Pendente', data_vencimento__lt=hoje).update(estado='Atrasado')")
    p(doc, "O dashboard admin atualiza pagamentos vencidos, calcula contagens e envia dados para graficos e cards.")

    h(doc, 2, "9.5 Padrao CRUD")
    three_col_table(
        doc,
        ["View", "Exemplo", "O que faz"],
        [
            ("list", "socio_list", "objects.all() e renderiza lista."),
            ("detail", "socio_detail", "get_object_or_404 pelo pk e mostra ficha."),
            ("create", "socio_create", "GET mostra form; POST cria."),
            ("update", "socio_update", "Passa instance ao form e guarda alteracoes."),
            ("delete", "socio_delete", "GET confirma; POST apaga."),
        ],
    )

    h(doc, 2, "9.6 Inscricoes sem duplicados")
    code(doc, "ja_existe = Inscricao.objects.filter(socio=socio_escolhido, aula=aula_escolhida).exists()")
    p(doc, "Antes de criar uma inscricao, o sistema confirma se o mesmo socio ja esta na mesma aula. Se estiver, mostra erro.")

    h(doc, 2, "9.7 Cancelamento seguro de inscricoes")
    code(doc, "if not request.user.is_superuser:\n    if not hasattr(request.user, 'socio') or inscricao.socio_id != request.user.socio.id:\n        return redirect('redirecionar_login')")
    p(doc, "Um socio so pode cancelar as suas proprias inscricoes. O admin pode gerir todas.")

    h(doc, 2, "9.8 Pagamentos e mensalidade automatica")
    code(doc, "if estado_anterior != 'Liquidado' and pagamento_guardado.estado == 'Liquidado' and pagamento_guardado.frequencia == 'Mensal':\n    Pagamento.objects.create(... estado='Pendente', frequencia='Mensal')")
    p(doc, "Quando uma mensalidade passa a Liquidado, o sistema cria automaticamente a proxima mensalidade pendente.")

    h(doc, 2, "9.9 Login e redirecionamento inteligente")
    code(doc, "if request.user.is_superuser:\n    return redirect('index')\nelif hasattr(request.user, 'treinador'):\n    return redirect('perfil_treinador')\nelse:\n    return redirect('perfil_cliente')")
    p(doc, "Apos login, o sistema manda cada tipo de utilizador para o sitio certo: admin, treinador ou socio.")

    h(doc, 2, "9.10 Registo de socio")
    code(doc, "user = User.objects.create_user(username=username, email=email, password=password)\nnovo_socio = Socio.objects.create(user=user, nome=nome, email=email, ...)\nPagamento.objects.create(socio=novo_socio, valor=35.00, estado='Pendente')\nlogin(request, user)")
    p(doc, "O registo cria um User, cria a ficha Socio ligada ao User, cria a primeira mensalidade e inicia sessao automaticamente.")

    h(doc, 2, "9.11 Perfil do cliente e QR Code")
    code(doc, "qr_data = f'SOCIO_ID:{socio.id}'\nqr = qrcode.QRCode(...)\nqr_base64 = base64.b64encode(buffer.getvalue()).decode()")
    p(doc, "O QR Code e gerado em memoria e enviado para o template em Base64. O perfil tambem mostra inscricoes, pagamentos, plano e aulas disponiveis.")

    h(doc, 2, "9.12 Perfil do treinador")
    p(doc, "Se request.user tiver treinador associado, a view mostra apenas as aulas desse treinador com Aulas.objects.filter(treinador=treinador).")

    h(doc, 2, "9.13 Central de Acessos")
    p(doc, "Apenas superusers entram. Permite criar admin, criar treinador com conta, associar conta a treinador ou socio existente, apagar conta de login e fazer reset password.")
    code(doc, "user_a_alterar.set_password(nova_password)\nuser_a_alterar.save()")
    p(doc, "set_password e essencial: guarda password com hash seguro, nao em texto simples.")

    h(doc, 1, "10. Templates")
    kv_table(
        doc,
        [
            ("base.html", "Base das paginas internas: navbar, Bootstrap, icons, CSS, favicon e estrutura comum."),
            ("landing.html", "Pagina publica, com carrosseis de modalidades e treinadores vindos da base de dados."),
            ("index.html", "Dashboard admin com estatisticas e graficos."),
            ("*_list.html", "Listagens."),
            ("*_detail.html", "Detalhes de um registo."),
            ("*_form.html", "Criar/editar com forms."),
            ("*_confirm_delete.html", "Confirmacao antes de apagar."),
            ("perfil_cliente.html", "Area do socio."),
            ("perfil_treinador.html", "Area do treinador."),
            ("gestao_acessos.html", "Central de contas e passwords."),
        ],
    )
    p(doc, "Nos templates usamos {% url 'nome_da_rota' %} para gerar links e {% static '...' %} para ficheiros estaticos. A view envia dados por context e o template usa variaveis e loops.")

    h(doc, 1, "11. admin.py")
    p(doc, "Regista os models no /admin/ do Django. list_display escolhe colunas visiveis, search_fields permite pesquisa e list_filter adiciona filtros.")

    h(doc, 1, "12. Deploy no Render")
    p(doc, "O GitHub guarda o codigo. O Render faz deploy automatico quando ha push. Durante o build instala dependencias, recolhe static files e aplica migracoes. Depois arranca com Gunicorn.")
    h(doc, 2, "12.1 Variaveis de ambiente")
    kv_table(
        doc,
        [
            ("SECRET_KEY", "Chave secreta do Django."),
            ("DEBUG", "False em producao."),
            ("DATABASE_URL", "Connection string do Neon/PostgreSQL."),
            ("ALLOWED_HOSTS", "powergym.site, www.powergym.site e .onrender.com."),
            ("CSRF_TRUSTED_ORIGINS", "Dominios HTTPS aceites para POST."),
            ("CLOUD_NAME / CLOUDINARY_CLOUD_NAME", "Nome da cloud Cloudinary."),
            ("API_KEY / CLOUDINARY_API_KEY", "Chave da API Cloudinary."),
            ("API_SECRET / CLOUDINARY_API_SECRET", "Segredo da API Cloudinary."),
        ],
    )
    h(doc, 2, "12.2 Comandos")
    code(doc, "Build Command:\npip install -r requirements.txt && python manage.py collectstatic --noinput && python manage.py migrate\n\nStart Command:\ngunicorn ginasio_project.wsgi:application")
    p(doc, "collectstatic prepara CSS/imagens/favicon. migrate aplica a estrutura na Neon. gunicorn corre a aplicacao Django em producao.")

    h(doc, 1, "13. Neon")
    p(doc, "A Neon e a base PostgreSQL em producao. O Django nao tem credenciais escritas no codigo: le DATABASE_URL no Render.")
    code(doc, "DATABASE_URL=postgresql://user:password@host/neondb?sslmode=require...")
    p(doc, "Na defesa nunca mostres a password real. Explica apenas que a string contem user, password, host, nome da base e SSL.")

    h(doc, 1, "14. Cloudinary")
    p(doc, "Cloudinary guarda imagens de socios e treinadores. Isto e melhor do que guardar uploads no disco do Render, porque o deploy pode reiniciar e nao devemos depender do filesystem local para ficheiros permanentes.")
    kv_table(
        doc,
        [
            ("ImageField", "Socio.foto e Treinador.foto."),
            ("request.FILES", "Recebe ficheiros enviados por formulario."),
            ("MediaCloudinaryStorage", "Envia uploads para Cloudinary."),
            ("foto.url", "URL usada nos templates para mostrar a imagem."),
        ],
    )

    h(doc, 1, "15. Static files, favicon e WhiteNoise")
    p(doc, "Ficheiros estaticos incluem CSS, logo, favicons e manifest. WhiteNoise permite servir esses ficheiros no Render.")
    code(doc, "path('favicon.ico', RedirectView.as_view(url=settings.STATIC_URL + 'images/favicon.ico', permanent=True))")
    p(doc, "Alguns browsers procuram /favicon.ico diretamente. Esta rota redireciona esse pedido para o favicon em /static/images/favicon.ico.")

    h(doc, 1, "16. requirements.txt")
    kv_table(
        doc,
        [
            ("Django", "Framework web principal."),
            ("gunicorn", "Servidor WSGI em producao."),
            ("dj-database-url", "Converte DATABASE_URL em configuracao Django."),
            ("psycopg2-binary", "Driver PostgreSQL para Neon."),
            ("whitenoise", "Static files em producao."),
            ("cloudinary/django-cloudinary-storage", "Uploads de imagens."),
            ("qrcode/pillow", "QR Code do socio."),
        ],
    )

    h(doc, 1, "17. Fluxos para treinar")
    flows = {
        "Registo de socio": [
            "Cliente abre /registo/.",
            "Preenche dados e submete.",
            "registo_socio verifica username/email.",
            "Cria User com create_user.",
            "Cria Socio ligado ao User.",
            "Cria pagamento inicial pendente.",
            "Faz login e redireciona para perfil_cliente.",
        ],
        "Login": [
            "LoginView valida credenciais.",
            "LOGIN_REDIRECT_URL chama redirecionar_login.",
            "Admin vai para dashboard.",
            "Treinador vai para perfil_treinador.",
            "Socio vai para perfil_cliente.",
        ],
        "Criar aula": [
            "Admin entra em /aulas/novo/.",
            "aula_create cria AulasForm.",
            "POST valida form.is_valid().",
            "form.save cria a aula.",
            "Redireciona para lista.",
        ],
        "Inscrever numa aula": [
            "Sistema recebe socio e aula.",
            "Verifica duplicado com exists().",
            "Se ja existe, mostra erro/aviso.",
            "Se nao existe, cria Inscricao.",
        ],
        "Reset password": [
            "Admin abre Central de Acessos.",
            "Seleciona utilizador e nova password.",
            "View chama set_password.",
            "Password fica com hash seguro.",
        ],
    }
    for title, steps in flows.items():
        h(doc, 2, title)
        for step in steps:
            number(doc, step)

    h(doc, 1, "18. Perguntas provaveis")
    qa = [
        ("Porque Django?", "Porque ja traz ORM, autenticacao, admin, forms, migrations, templates e seguranca base."),
        ("O que e ORM?", "E a camada que permite trabalhar com tabelas atraves de classes Python, como Socio.objects.filter(...)."),
        ("Onde esta a base de dados?", "Em producao esta na Neon/PostgreSQL, ligada pelo DATABASE_URL no Render."),
        ("Porque Cloudinary?", "Para guardar uploads de imagens fora do Render, de forma permanente."),
        ("Como protegem admin?", "Com admin_required e request.user.is_superuser."),
        ("Como distinguem roles?", "Admin e is_superuser; treinador tem request.user.treinador; socio tem request.user.socio."),
        ("Password em texto simples?", "Nao. create_user e set_password guardam hashes."),
        ("Para que serve migrate?", "Aplica as migrations na base de dados."),
        ("Para que serve collectstatic?", "Prepara static files para producao."),
    ]
    for q, a in qa:
        p(doc, "Pergunta: " + q)
        bullet(doc, "Resposta: " + a)

    h(doc, 1, "19. Melhorias finais feitas no projeto")
    for item in [
        "Landing page reformulada.",
        "Carrosseis de modalidades e treinadores carregados da base de dados.",
        "Paginas internas mais profissionais.",
        "Mensagem de erro no login invalido.",
        "Central de Acessos com reset password.",
        "Favicon e manifest corrigidos.",
        "Configuracao Render/Neon/Cloudinary por variaveis de ambiente.",
    ]:
        bullet(doc, item)

    h(doc, 1, "20. Resumo final para memorizar")
    p(
        doc,
        "O PowerGym e uma aplicacao Django completa para gestao de ginasio. Os models representam as tabelas da base de dados e as relacoes. As views aplicam a logica, protegem acessos, usam forms e enviam dados para templates. A base de dados de producao e PostgreSQL na Neon, ligada por DATABASE_URL no Render. O Render faz deploy a partir do GitHub com collectstatic, migrate e gunicorn. O Cloudinary guarda imagens. A autenticacao usa o User do Django e separa admin, socio e treinador.",
    )


def main():
    doc = setup_doc()
    add_cover(doc)
    add_content(doc)
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("PowerGym - Guia de Defesa Tecnica")
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED
    if OUT.exists():
        OUT.unlink()
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
